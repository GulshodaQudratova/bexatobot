from aiogram import types
from loader import dp
from aiogram.types import InputFile
from docx import Document
from function import to_latin,to_cyrillic
import io
from io import BytesIO
@dp.message_handler(content_types=types.ContentTypes.DOCUMENT)
async def test(message:types.Message):
    document_name= message.document.file_name
    from main import has_cyrillic
    file = BytesIO()
    await message.document.download(file)
    file.seek(0)
    try:
       document = Document(file)
       for para in document.paragraphs:
           if has_cyrillic(para.text):
               para.text = to_latin(para.text)
           else:
               para.text =  to_cyrillic(para.text)

       for table in document.tables:
           for row in table.rows:
               for cell in row.cells:
                   if has_cyrillic(cell.text):
                       cell.text = to_latin(cell.text)
                   else:
                       cell.text = to_cyrillic(cell.text)
       document2 = BytesIO()
       document.save(document2)
       document2.seek(0)
       send_file = InputFile(document2,f'{document_name}')
       await message.answer_document(document=send_file)
    except Exception as e:
        print(e)
        await message.answer("Iltimos Word tipidagi(docx tipida) yuboring!")